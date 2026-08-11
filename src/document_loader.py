"""
Ingestion: upload -> extract -> clean -> split.

Supported formats: PDF, DOCX, TXT, MD.
Each output Chunk carries enough metadata (source file, page number, chunk
index) to be cited back to the user later in the "Sources" panel.
"""

from __future__ import annotations

import re
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

import pypdf
import docx

from src.text_splitter import RecursiveCharacterTextSplitter


class UnsupportedFileType(ValueError):
    pass


@dataclass
class Chunk:
    text: str
    source: str
    chunk_index: int
    page: Optional[int] = None
    doc_id: str = ""

    @property
    def id(self) -> str:
        return f"{self.doc_id}:{self.chunk_index}"


# --------------------------------------------------------------------------- #
# Extraction
# --------------------------------------------------------------------------- #
def extract_pdf(path: str) -> List[Tuple[Optional[int], str]]:
    """Returns a list of (page_number, raw_text) tuples, 1-indexed pages."""
    reader = pypdf.PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append((i + 1, text))
    return pages


def extract_docx(path: str) -> List[Tuple[Optional[int], str]]:
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return [(None, "\n".join(parts))]


def extract_plain_text(path: str) -> List[Tuple[Optional[int], str]]:
    text = Path(path).read_text(encoding="utf-8", errors="ignore")
    return [(None, text)]


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_plain_text,
    ".md": extract_plain_text,
}


def extract(path: str) -> List[Tuple[Optional[int], str]]:
    ext = Path(path).suffix.lower()
    if ext not in EXTRACTORS:
        raise UnsupportedFileType(
            f"'{ext}' is not supported. Supported types: {', '.join(EXTRACTORS)}"
        )
    return EXTRACTORS[ext](path)


# --------------------------------------------------------------------------- #
# Cleaning
# --------------------------------------------------------------------------- #
def clean_text(text: str) -> str:
    if not text:
        return ""
    # Re-join words that were hyphenated across a line break: "exam-\nple".
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Collapse runs of spaces/tabs.
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ newlines down to a paragraph break.
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim stray spaces around newlines.
    text = re.sub(r" *\n *", "\n", text)
    # Drop common PDF artefacts: lone page-number lines, repeated form-feeds.
    text = re.sub(r"\n\s*\d{1,4}\s*\n", "\n", text)
    text = text.replace("\x0c", "\n")
    return text.strip()


# --------------------------------------------------------------------------- #
# Full pipeline: file -> cleaned, chunked, metadata-tagged pieces
# --------------------------------------------------------------------------- #
def load_and_chunk(
    path: str,
    display_name: Optional[str] = None,
    chunk_size: int = 1000,
    chunk_overlap: int = 150,
) -> List[Chunk]:
    display_name = display_name or Path(path).name
    doc_id = uuid.uuid4().hex[:12]
    splitter = RecursiveCharacterTextSplitter(chunk_size, chunk_overlap)

    pages = extract(path)
    chunks: List[Chunk] = []
    running_index = 0
    for page_num, raw_text in pages:
        cleaned = clean_text(raw_text)
        if not cleaned:
            continue
        for piece in splitter.split_text(cleaned):
            chunks.append(
                Chunk(
                    text=piece,
                    source=display_name,
                    chunk_index=running_index,
                    page=page_num,
                    doc_id=doc_id,
                )
            )
            running_index += 1
    return chunks
